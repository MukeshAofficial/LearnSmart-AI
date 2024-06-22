from flask import Flask, render_template, jsonify, request,  send_file,  redirect, url_for
import json
import io
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import requests
from IPython.display import display, Markdown
import google.generativeai as genai
import base64
from io import BytesIO
import io
import os
import PIL.Image
from PIL import Image
import textwrap
from docx import Document
from googletrans import Translator

app = Flask(__name__)

translator = Translator()

API_KEY = 'Api-Key' # YOUR CRACKEDDEVS API KEY HERE
LIMIT = 10  # Number of jobs to fetch

def get_jobs():
    url = f'https://api.crackeddevs.com/v1/get-jobs?limit={LIMIT}'
    headers = {
        'api-key': API_KEY,
    }

    response = requests.get(url, headers=headers)

    if response.ok:
        jobs = response.json()
        for job in jobs:
            job['description'] = convert_markdown_to_html(job['description'])
        return jobs
    else:
        return []

def convert_markdown_to_html(markdown_text):
    # Replace markdown elements with HTML
    html_text = markdown_text.replace('**', '<strong>').replace('*', '<li>')
    html_text = html_text.replace('\n\n', '</p><p>').replace('\n', '<br>')
    html_text = '<p>' + html_text + '</p>'
    return html_text

genai.configure(api_key="Apit-KEy")

model = genai.GenerativeModel('gemini-pro-vision')


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/job')
def job():
    jobs = get_jobs()
    return render_template('jobs.html', jobs=jobs)


def to_markdown(text):
    text = text.replace('\u2022', '  *')
    return Markdown(textwrap.indent(text, '> ', predicate=lambda _: True))



@app.route('/resume')
def resume():
    return render_template('resume.html')



@app.route('/upload', methods=['POST'])
def upload():
    try:
        # Get the image file from the request
        resume_file = request.files['resume']

        # Save the image to a file on your server
        img_path = 'resume.jpeg'
        resume_file.save(img_path)

        # Load the image
        img = PIL.Image.open(img_path)

        # Generate text from the image using Generative AI model
        response = model.generate_content(["Please review my resume for improvements, additions, and mistakes. Provide a score out of 10 based on clarity, structure, relevance to job roles, readability, formatting, and ATS keyword optimization.", img], stream=True)
        response.resolve()

        # Store response text in a variable
        
        review_text = response.text
        notepad_file_path = 'review.txt'
        with open(notepad_file_path, 'w') as file:
            file.write(response.text)


        # Render the review.html template with the review text
        return render_template('review.html', review_text=review_text)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

def strip_code(response_text):
    response_text = response_text.strip()
    if response_text.startswith('```json') and response_text.endswith('```'):
        return response_text[7:-3].strip()
    return response_text

def generate_certificate_pdf(name, certificate_text, template_path, output_path):
    template_pdf = PyPDF2.PdfReader(template_path)
    template_page = template_pdf.pages[0]

    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)

    can.setFont("Helvetica-Bold", 24)
    name_width = can.stringWidth(name, "Helvetica-Bold", 24)
    x = ((letter[0] - name_width) / 2) + 100
    y = letter[1] - 500
    can.drawString(x, y, name)

    can.setFont("Helvetica", 12)
    text_width = can.stringWidth(certificate_text, "Helvetica", 12)
    x = ((letter[0] - text_width) / 2) + 100
    y -= 30
    can.drawString(x, y, certificate_text)

    can.save()
    packet.seek(0)

    new_pdf = PyPDF2.PdfReader(packet)
    output = PyPDF2.PdfWriter()
    template_page.merge_page(new_pdf.pages[0])
    output.add_page(template_page)

    with open(output_path, 'wb') as f:
        output.write(f)


@app.route('/quiz', methods=['GET', 'POST'])
def quiz():
    topic = ""
    questions = []

    if request.method == 'POST':
        topic = request.form.get('topic')
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(
            f"""Create a 5-question multiple-choice quiz on the topic of {topic}. 
               Provide the questions, four answer options for each question, and the correct answer. 
               Format the output as a JSON array with each question as an object containing: 
               'question', 'options' (list of four strings), and 'answer' (string)."""
        )
        response_text = response.text
        stripped_text = strip_code(response_text)

        try:
            questions = json.loads(stripped_text)
        except json.JSONDecodeError as e:
            print(f"Error parsing response: {e}")
            questions = []

    return render_template('quiz.html', topic=topic, questions=questions)

@app.route('/submit', methods=['POST'])
def submit():
    questions = json.loads(request.form.get('questions'))
    user_answers = request.form.to_dict(flat=True)
    user_answers.pop('questions')

    correct_count = 0
    total_questions = len(questions)

    for i, question in enumerate(questions):
        correct_answer = question['answer']
        user_answer = user_answers.get(f'question-{i}')
        if user_answer == correct_answer:
            correct_count += 1

    return render_template('results.html', total=total_questions, correct=correct_count, incorrect=total_questions - correct_count)

@app.route('/certificate', methods=['GET', 'POST'])
def certificate():
    if request.method == 'POST':
        name = request.form.get('username')
        certificate_text = f"{name} has successfully completed the quiz."
        template_path = 'certificate.pdf'
        output_path = 'certificate_.pdf'

        generate_certificate_pdf(name, certificate_text, template_path, output_path)
        return send_file(output_path, as_attachment=True)

    return render_template('certificate.html')

@app.route('/codemaster')
def code():
    return render_template('code.html')

@app.route('/generate', methods=['POST'])
def generate_html():
    prompt = request.form['prompt']
    model = genai.GenerativeModel('gemini-1.5-flash')
    rply = model.generate_content("Generate both html and css code in single file for " + prompt + " with colorful CSS background and more attractive CSS, I need only code and no explanation")
    html_content = rply.text

    with open("templates/index1.html", "w") as file:
        file.write(html_content)

    return redirect(url_for('output'))


@app.route('/output')
def output():
    return render_template('index1.html')

@app.route('/chatbot')
def chatbot():
    return render_template('chatbot.html')

@app.route('/get_response', methods=['POST'])
def get_response():
    data = request.json
    user_message = data['message']
    model = genai.GenerativeModel('gemini-1.5-flash')
    rply = model.generate_content(f"{user_message}  answer in one or 2 lines without any */ symbols")

    return jsonify({'response': rply.text})

@app.route('/to-do-list')
def to_do_list():
    return render_template('to-do-list.html')

@app.route('/timer')
def timer():
    return render_template('timer.html')


@app.route('/speech')
def speechindex():
    return render_template('speech.html')

@app.route('/store_transcription', methods=['POST'])
def store_transcription():
    data = request.json
    transcription = data['transcription']

    # Store transcription in a Word file
    document = Document()
    document.add_paragraph(transcription)
    document.save('transcription.docx')

    return 'Transcription stored successfully'

@app.route('/speech-summary')
def speech():
    with open('transcription.docx', 'rb') as docx_file:
        doc = Document(docx_file)
        script_summary = ""
        for para in doc.paragraphs:
            script_summary += para.text + "<br>"
    
    # Render HTML template with the content
    return render_template('transcript-result.html', script_summary=script_summary)




@app.route('/translator')
def translator_page():
    return render_template('translator.html')

@app.route('/translate', methods=['POST'])
def translate():
    text = request.form['text']
    source_lang = request.form['source_lang']
    target_lang = request.form['target_lang']
    
    translated_text = translator.translate(text, src=source_lang, dest=target_lang).text
    return jsonify({'translated_text': translated_text})


@app.route('/draw')
def draw():
    return render_template('draw-and-guess.html')

@app.route('/guess', methods=['POST','GET'])
def guess():
    try:
        data = request.get_json()
        img_data = data['image']
        img_data = img_data.split(',')[1]
        img_data = base64.b64decode(img_data)
        img = Image.open(BytesIO(img_data))
        img.save('draw.png')

        response = model.generate_content(["Guess my drawing .", img], stream=True)
        response.resolve()

        # Store response text in a variable
        guess = response.text
        print(guess)
        
        return render_template('guess.html', guess=guess)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500  # Handle errors gracefully
 

@app.route('/scan-solve')
def scanindex():
    return render_template('scan-solve.html')

@app.route('/upload1', methods=['POST','GET'])
def upload2():
    try:
        # Get the image data from the request
        data = request.json
        image_data = data.get('image')

        # Decode the base64 image data
        image_data = image_data.split(',')[1]
        image_binary = base64.b64decode(image_data)

        # Save the image to a file on your local PC
        img_path = 'question.png'
        with open(img_path, 'wb') as img_file:
            img_file.write(image_binary)

        # Load the image
        img = PIL.Image.open(io.BytesIO(image_binary))

        # Use Generative AI model to generate text from the image
        response = model.generate_content(["Give me a simple answer for this question", img], stream=True)
        response.resolve()

        # Store response text in a Notepad file
        notepad_file_path = 'solution.txt'
        with open(notepad_file_path, 'w') as file:
            file.write(response.text)
            file_written = True
        
        if file_written:
            return render_template('scan-result.html', solution=response.text)
        else:
            return render_template('error.html', message="Failed to write response to file.")

        return render_template('scan-result.html', solution=response.text)
    

    except Exception as e:
        return jsonify({'error': str(e)}), 500 


def generate_notesummary(note_text):
    model = genai.GenerativeModel('gemini-pro')
    rply = model.generate_content("summarize my notes"+note_text)
    to_markdown(rply.text)
    return rply.text

@app.route('/note', methods=['GET', 'POST'])
def note():
    if request.method == 'POST':
        # Check if a file was uploaded
        if 'Note_file' not in request.files:
            return render_template('error.html', message='No file part')

        file = request.files['Note_file']

        # Check if the file is empty
        if file.filename == '':
            return render_template('error.html', message='No selected file')

        # Check if the file is of allowed type
        if file and file.filename.endswith('.txt'):
            # Read the file content
            note_text = file.read().decode('utf-8')

            # Generate summary
            summary_text = generate_notesummary(note_text)
            
            # Render the result template with summary
            return render_template('note-result.html', summary_text=summary_text)

        else:
            return render_template('error.html', message='Invalid file type. Please upload a text file')

    return render_template('note.html')

if __name__ == '__main__':
    app.run(debug=True)
